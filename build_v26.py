# -*- coding: utf-8 -*-
"""Build MOPLUS_SRS_EFF_WEB_v2.6.docx = v2.5 + Potable Water module (11 functions)
appended at end of chapter 3, from SRS-MO-Water-Web-v1.0.md, with embedded mockup
screenshots, changelog row, version bump. Does NOT overwrite v2.5."""
import re, io, os
import docx
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC_DOCX = r'ba/workspace/input/Customer_docs/documents/MOPLUS_SRS_EFF_WEB_v2.5.docx'
OUT_DOCX = r'ba/workspace/input/Customer_docs/documents/MOPLUS_SRS_EFF_WEB_v2.6.docx'
MD = r'ba/workspace/drafts/srs/03-dac-ta-chuc-nang/SRS-MO-Water-Web-v1.0.md'
IMGDIR = r'ba/workspace/drafts/mockup/srs-img'
IMG_MAP = {'YC33':'water.png','YC34':'yc34.png','YC35':'yc35.png','YC36':'yc36.png',
           'YC37':'yc37.png','YC38':'yc38.png','YC25':'yc25.png','YC26':'yc26.png',
           'YC27':'yc27.png','YC28':'yc28.png'}

# ---------- docx helpers ----------
def set_borders(table):
    tblPr=table._tbl.tblPr
    b=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000')
        b.append(e)
    tblPr.append(b)

def shade(cell, fill='D9D9D9'):
    tcPr=cell._tc.get_or_add_tcPr()
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),fill)
    tcPr.append(s)

def cell_text(cell, text, bold=False):
    cell.text=''
    p=cell.paragraphs[0]
    # support **bold** inside cell
    parts=text.split('**')
    if len(parts)==1:
        r=p.add_run(text); r.bold=bold
    else:
        for i,seg in enumerate(parts):
            if seg=='' : continue
            r=p.add_run(seg); r.bold=(bold or (i%2==1))
    for r in p.runs:
        r.font.name='Times New Roman'; r.font.size=Pt(11)

def add_table(doc, rows, header=True, summary=False, widths=None):
    ncol=len(rows[0])
    t=doc.add_table(rows=0, cols=ncol)
    set_borders(t)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(cells):
            txt=row[ci] if ci<len(row) else ''
            if summary:
                cell_text(val, txt, bold=(ci==0))
            elif header and ri==0:
                cell_text(val, txt, bold=True); shade(val)
            else:
                cell_text(val, txt, bold=False)
    if widths:
        for ci,w in enumerate(widths):
            for r in t.rows:
                r.cells[ci].width=Cm(w)
    return t

def add_para(doc, text, italic=False):
    p=doc.add_paragraph()
    parts=text.split('**')
    for i,seg in enumerate(parts):
        if seg=='' : continue
        r=p.add_run(seg)
        r.bold=(i%2==1)
        r.italic=italic
        r.font.name='Times New Roman'; r.font.size=Pt(12)
    return p

def add_image(doc, path):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run()
    run.add_picture(path, width=Cm(16.5))

# ---------- markdown parse ----------
with io.open(MD,'r',encoding='utf-8') as f:
    lines=f.read().split('\n')

# strip frontmatter
if lines[0].strip()=='---':
    end=lines.index('---',1); lines=lines[end+1:]

def is_table_line(l): return l.strip().startswith('|')
def is_sep(cells): return all(re.match(r'^:?-{2,}:?$', c.strip()) for c in cells if c.strip()!='' ) and any(c.strip() for c in cells)
def parse_cells(l):
    return [c.strip() for c in l.strip().strip('|').split('|')]

ops=[]  # list of ('type', payload)
i=0
intro_done=False
skip_changelog_section=False
FUNC_RE=re.compile(r'^##\s+[ABC]\d+\.\s+(.*)')
while i < len(lines):
    raw=lines[i]; l=raw.strip()
    if l=='' :
        i+=1; continue
    # tables
    if is_table_line(raw):
        block=[]
        while i<len(lines) and is_table_line(lines[i]):
            cs=parse_cells(lines[i])
            if not is_sep(cs): block.append(cs)
            i+=1
        if block:
            ops.append(('table', block))
        continue
    # H1 title -> skip
    if l.startswith('# ') and not l.startswith('## '):
        i+=1; continue
    m=FUNC_RE.match(l)
    if l.startswith('## '):
        title=l[3:].strip()
        if m:
            ops.append(('h2', m.group(1).strip()))
        elif title.startswith('NHÓM'):
            pass  # skip group heading, keep following note paragraph
        elif title.startswith('Câu hỏi'):
            ops.append(('h2','Nước sạch — Câu hỏi / Nội dung cần làm rõ'))
        elif title.startswith('Dòng cập nhật Change Log'):
            skip_changelog_section=True  # skip until next ## / ---
        elif title.startswith('Ma trận truy vết'):
            skip_changelog_section=False
            ops.append(('h2','Nước sạch — Ma trận truy vết'))
        else:
            ops.append(('h2', title))
        i+=1; continue
    if l.startswith('### '):
        if not skip_changelog_section:
            ops.append(('h3', l[4:].strip()))
        i+=1; continue
    if l.startswith('---'):
        skip_changelog_section=False
        i+=1; continue
    if l.startswith('> '):
        mm=re.search(r'\[ẢNH:\s*(YC\d+)\]', l)
        if mm:
            ops.append(('image', mm.group(1)))
        i+=1; continue
    # plain paragraph (may be italic *...*)
    if skip_changelog_section:
        i+=1; continue
    italic=False; txt=l
    if l.startswith('*') and l.endswith('*') and not l.startswith('**'):
        italic=True; txt=l.strip('*').strip()
    ops.append(('para', txt, italic))
    i+=1

# ---------- build docx ----------
doc=docx.Document(SRC_DOCX)

# lead-in heading for the module (H2, flat like other ch.3 functions)
h=doc.add_heading('', level=2)
r=h.add_run('Phân hệ Quản lý nước sạch (Potable Water Service)')

img_missing=[]
for op in ops:
    t=op[0]
    if t=='h2':
        doc.add_heading(op[1], level=2)
    elif t=='h3':
        doc.add_heading(op[1], level=3)
    elif t=='para':
        add_para(doc, op[1], italic=op[2])
    elif t=='image':
        fn=IMG_MAP.get(op[1])
        path=os.path.join(IMGDIR, fn) if fn else None
        if path and os.path.exists(path):
            add_image(doc, path)
        else:
            img_missing.append(op[1]); add_para(doc, '[Ảnh màn hình: %s]'%op[1], italic=True)
    elif t=='table':
        rows=op[1]
        ncol=len(rows[0])
        summary = (ncol==2 and rows[0][0].startswith('Tên chức năng'))
        add_table(doc, rows, header=(not summary), summary=summary)
    # spacing after tables/images
    if t in ('table','image'):
        doc.add_paragraph('')

# ---------- changelog row ----------
cl=doc.tables[2]
desc=('Bổ sung phân hệ Quản lý nước sạch (Potable Water Service): Tab xem/nhập tay chỉ số '
      'nước chuyến bay, 5 báo cáo đối chiếu & giám sát, 4 màn cấu hình tham số/danh mục, xuất Excel.')
vals=['18/06/2026','Mục Nước sạch (cuối chương 3)','A',desc,'2.6']
# find first empty data row
target=None
for r in cl.rows[1:]:
    if all(c.text.strip()=='' for c in r.cells):
        target=r; break
if target is None:
    target=cl.add_row()
for ci,c in enumerate(target.cells):
    cell_text(c, vals[ci] if ci<len(vals) else '')

# ---------- version bump on cover ----------
for tb in doc.tables:
    for row in tb.rows:
        for c in row.cells:
            if 'MOPLUS_SRS_EFF_WEB_v2.5' in c.text:
                cell_text(c, c.text.replace('v2.5','v2.6'))

doc.save(OUT_DOCX)
print('SAVED', OUT_DOCX)
print('ops:', len(ops), '| missing images:', img_missing)
